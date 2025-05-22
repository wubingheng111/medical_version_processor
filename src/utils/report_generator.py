import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import io
from PIL import Image

class ReportGenerator:
    def __init__(self, output_dir="reports"):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def save_analysis_report(self, original_image, processed_image, analysis_text, operation_type, params=None):
        """生成分析报告并保存为Word文档"""
        doc = Document()
        
        # 添加标题
        doc.add_heading('医学图像处理分析报告', 0)
        
        # 添加基本信息
        doc.add_heading('基本信息', level=1)
        doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'处理类型：{operation_type}')
        
        if params:
            doc.add_heading('处理参数', level=1)
            for key, value in params.items():
                doc.add_paragraph(f'{key}: {value}')
        
        # 添加图像
        doc.add_heading('图像对比', level=1)
        
        # 保存原始图像
        orig_img_path = self._save_temp_image(original_image, "original.png")
        doc.add_picture(orig_img_path, width=Inches(3.0))
        doc.add_paragraph('原始图像')
        
        # 保存处理后图像
        proc_img_path = self._save_temp_image(processed_image, "processed.png")
        doc.add_picture(proc_img_path, width=Inches(3.0))
        doc.add_paragraph('处理后图像')
        
        # 添加分析结果
        doc.add_heading('AI分析结果', level=1)
        doc.add_paragraph(analysis_text)
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"analysis_report_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        # 清理临时文件
        os.remove(orig_img_path)
        os.remove(proc_img_path)
        
        return filepath

    def _save_temp_image(self, image, filename):
        """保存临时图像文件"""
        if isinstance(image, Image.Image):
            img = image
        else:
            img = Image.fromarray(image)
        
        filepath = os.path.join(self.output_dir, filename)
        img.save(filepath)
        return filepath

    def generate_batch_export_report(self, export_results, export_path):
        """生成批量导出报告
        
        Args:
            export_results: 导出结果列表
            export_path: 导出目录路径
        """
        doc = Document()
        
        # 添加标题
        doc.add_heading('医学图像批处理导出报告', 0)
        
        # 添加基本信息
        doc.add_heading('导出信息', level=1)
        doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'导出路径：{export_path}')
        doc.add_paragraph(f'总文件数：{len(export_results)}')
        
        # 添加处理统计
        successful = sum(1 for r in export_results if r.get('success', False))
        failed = len(export_results) - successful
        
        doc.add_heading('处理统计', level=1)
        doc.add_paragraph(f'成功处理：{successful}')
        doc.add_paragraph(f'处理失败：{failed}')
        
        # 添加详细信息
        doc.add_heading('处理详情', level=1)
        
        # 成功处理的文件
        doc.add_heading('成功处理的文件', level=2)
        for result in (r for r in export_results if r.get('success', False)):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"文件名：{result.get('file_name', 'unknown')}")
            if 'type' in result:
                p.add_run(f"\n处理类型：{result['type']}")
            if 'quality_metrics' in result:
                p.add_run("\n图像质量指标：")
                for metric, value in result['quality_metrics'].items():
                    p.add_run(f"\n  - {metric}: {value:.2f}")
        
        # 处理失败的文件
        if failed > 0:
            doc.add_heading('处理失败的文件', level=2)
            for result in (r for r in export_results if not r.get('success', False)):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"文件名：{result.get('file_name', 'unknown')}")
                if 'error' in result:
                    p.add_run(f"\n失败原因：{result['error']}")
        
        # 保存报告
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"batch_export_report_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath 